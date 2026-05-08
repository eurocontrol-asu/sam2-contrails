"""Generate the SAM2 labelling guide DOCX — Edition 2.0.

Produces a clear, step-by-step guide for labellers reviewing SAM2 contrail
predictions in Encord.  Written for non-experts.

Key design decisions:
  - VIDEO-first: emphasises temporal navigation as the primary review tool
  - Prompt-by-Prompt workflow: each flight is reviewed independently
  - Colour convention matches Encord: Prompt=orange, Contrail=red
  - Every instruction is concrete and unambiguous

Image placeholders (marked [INSERT FIGURE ...]) are left for the user to
insert their own annotated screenshots from Encord.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_PATH = Path(__file__).parent / "labelling_guide_sam2.docx"


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D6E4F0")

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_important(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Important:  ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    p.add_run(text)


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Tip:  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    p.add_run(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT FIGURE: {text}]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(150, 0, 150)


def add_caption(doc, text):
    cap = doc.add_paragraph(text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Labelling Guide")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SAM2 Contrail Review Campaign")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")
    doc.add_paragraph("")

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(
        "A step-by-step guide for reviewing contrail predictions\n"
        "made by the SAM2 model on ground-camera video sequences."
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Edition 2.0  |  May 2026  |  EUROCONTROL")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  What is this campaign about?",
        "2.  The data: video sequences",
        "3.  Key concepts: Prompts and Contrails",
        "    3.1  Prompt (orange polygon)",
        "    3.2  Contrail (red multi-polygon)",
        "    3.3  How Prompts and Contrails relate spatially",
        "    3.4  Confidence score",
        "    3.5  Out-of-scope contrails (no Prompt = do not label)",
        "4.  How the model works",
        "5.  Using temporal navigation",
        "6.  Your tasks as a labeller (Prompt by Prompt)",
        "    6.1  For each Prompt: check if a real contrail is visible",
        "    6.2  Verify the Contrail-to-Prompt link",
        "    6.3  Fix polygon shapes",
        "    6.4  Delete false positives (sun artifacts)",
        "7.  Common mistakes to avoid",
        "8.  Step-by-step Encord workflow",
        "9.  Quick reference checklist",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. WHAT IS THIS CAMPAIGN ABOUT?
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("1  What is this campaign about?", level=1)

    doc.add_paragraph(
        "Aircraft engines produce condensation trails (also known as contrails) "
        "that are visible as white lines in the sky. EUROCONTROL operates a "
        "ground-based sky camera in "
        u"Brétigny-Sur-Orge (south of Paris) "
        "that photographs the sky every 30 seconds, all day long. "
        "These images are used to study contrails and their environmental impact."
    )
    doc.add_paragraph(
        "A machine-learning model called SAM2 (Segment Anything 2) has been "
        "trained to automatically detect and track contrails in these video "
        "sequences. The model is good, but not perfect. Your job is to review "
        "the model's predictions and correct any mistakes."
    )

    doc.add_heading("What has changed from the previous campaign?", level=2)
    doc.add_paragraph(
        "In the previous campaign, you drew contrail polygons from scratch. "
        "This time is different:"
    )
    add_bullets(doc, [
        "The SAM2 model has already drawn polygons around the contrails it detected.",
        "Each polygon is linked to the flight that created the contrail.",
        "You review and correct these polygons rather than drawing them from scratch.",
        "This is faster, but you need to watch for specific types of model errors.",
    ])

    add_important(doc,
        "You are not starting from scratch. The polygons are already there. "
        "Your role is quality control: verify, correct, and complete."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. THE DATA: VIDEO SEQUENCES
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("2  The data: video sequences", level=1)

    doc.add_paragraph(
        "Each labelling task is a video sequence covering a 2-hour window, "
        "containing up to 240 frames (one frame every 30 seconds). The images "
        "are 1024 x 1024 pixels, showing the sky from the camera's rooftop "
        "location."
    )

    add_important(doc,
        "You are reviewing videos, not static images. Navigating forward "
        "and backward through the frames is your most powerful tool. A single "
        "frame can be ambiguous; watching how a feature evolves over time "
        "almost always resolves the ambiguity."
    )

    doc.add_paragraph(
        "The appearance of the sky changes throughout the day:"
    )
    add_bullets(doc, [
        "Morning: the sun may cause bright glare in parts of the image.",
        "Midday: clearest visibility, best conditions for seeing contrails.",
        "Afternoon/evening: sun glare may appear again from a different direction.",
        "Cloudy days: contrails may be partially hidden behind or above clouds.",
    ])

    doc.add_paragraph(
        "As you step through the video, you will see contrails appear, grow, "
        "drift with the wind, and eventually fade away. Prompts (flight "
        "trajectories) will appear and disappear as different flights pass "
        "through the camera's field of view."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. KEY CONCEPTS
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("3  Key concepts: Prompts and Contrails", level=1)

    doc.add_paragraph(
        "Every annotation in this campaign belongs to one of two categories: "
        "Prompt or Contrail. They are visually distinct in Encord by colour."
    )

    # ── 3.1 Prompt ──
    doc.add_heading("3.1  Prompt (orange polygon)", level=2)
    doc.add_paragraph(
        "A Prompt is an orange polygon that shows where a specific flight "
        "passed overhead during the last 5 minutes of its trajectory above "
        "the camera. It is computed from the aircraft's actual flight path "
        "and wind data (advection): the system projects the trajectory onto "
        "the camera image and shifts it according to wind speed and direction "
        "to account for where the contrail would have drifted."
    )
    doc.add_paragraph(
        "Because Prompts only cover the last 5 minutes of the flight, they "
        "indicate where a recent contrail should be — if one was produced. "
        "Not every flight produces a visible contrail."
    )
    doc.add_paragraph(
        "Each Prompt carries a flight_id attribute — a text label identifying "
        'the flight (for example, "SWR15X_1161").'
    )

    add_important(doc,
        "Prompts are always orange polygons. Never delete or modify a Prompt. "
        "They are part of the input data and represent known flight trajectories."
    )

    # ── 3.2 Contrail ──
    doc.add_heading("3.2  Contrail (red multi-polygon)", level=2)
    doc.add_paragraph(
        "A Contrail is a red multi-polygon drawn by the model around where "
        "it thinks a contrail actually exists in the image. It is the model's "
        "prediction — what you are reviewing."
    )
    doc.add_paragraph(
        "Contrails vary in shape depending on their age and atmospheric "
        "conditions:"
    )
    add_bullets(doc, [
        "Young contrails (< 5 minutes old) are typically elongated and thin, "
        "closely following the aircraft's flight path.",
        "Older contrails may spread wider and develop more rounded, irregular "
        "shapes as turbulence and wind shear distort them over time.",
        "Very old contrails can become diffuse and cloud-like — at that point "
        "they should no longer be labelled.",
    ])

    doc.add_paragraph(
        "Each Contrail carries two attributes:"
    )
    add_bullets(doc, [
        "flight #relation — links this Contrail to the Prompt (flight) that "
        "created it. This link is critical: it tells us which flight produced "
        "which contrail.",
        "score — the model's confidence (between 0 and 1) that this detection "
        "is a real contrail. High scores (above 0.80) are usually correct. "
        "Low scores (around 0.50) are often false positives.",
    ])

    doc.add_paragraph("")
    add_placeholder(doc, "Ontology screenshot — Prompt (orange) and Contrail (red) in Encord")
    add_caption(doc,
        "Figure 1 — The two object types. In Encord, Prompts are orange, "
        "Contrails are red. Every Contrail must be linked to exactly one Prompt."
    )

    doc.add_paragraph("")
    add_styled_table(doc,
        headers=["Object", "Colour in Encord", "Attributes", "What it represents"],
        rows=[
            ["Prompt", "Orange", "flight_id (text)",
             "Where a flight passed in the last 5 min (wind-advected). "
             "Input data — never edit or delete."],
            ["Contrail", "Red", "flight #relation, score",
             "Model's prediction of a contrail. Must be linked to the "
             "Prompt of the flight that created it."],
        ],
        col_widths=[2.5, 2.5, 3, 7],
    )

    # ── 3.3 Spatial relationship ──
    doc.add_paragraph("")
    doc.add_heading("3.3  How Prompts and Contrails relate spatially", level=2)

    doc.add_paragraph(
        "A Contrail should be aligned with its linked Prompt — both should "
        "follow the same general direction in the image, because the contrail "
        "was produced by that flight. However, they do not necessarily overlap:"
    )
    add_bullets(doc, [
        "Young contrails (< 5 min old) will typically overlap with their "
        "Prompt, because the Prompt covers the most recent 5 minutes of "
        "the flight trajectory and the contrail has not had time to drift far.",
        "Older contrails (> 5 min old) may extend well beyond the Prompt, "
        "because the contrail has survived longer than the 5-minute window "
        "the Prompt covers. The Contrail and Prompt will still be aligned in "
        "direction, but the Contrail may be spatially offset, wider, or "
        "longer than the Prompt.",
    ])

    add_important(doc,
        "\"Aligned\" means the Contrail follows the same direction as the "
        "Prompt, not that they must overlap pixel-for-pixel. Think of the "
        "Prompt as a directional guide showing where the flight passed, not "
        "as an exact boundary of the contrail."
    )

    # ── 3.4 Score ──
    doc.add_heading("3.4  Confidence score", level=2)
    doc.add_paragraph(
        "Every Contrail carries a score attribute — the model's confidence "
        "that the detected feature is a real contrail. The score is a number "
        "between 0 and 1, visible in the attributes panel when you click on "
        "a Contrail in Encord."
    )
    add_bullets(doc, [
        "Scores above 0.80 — the model is confident. These are almost always "
        "real contrails. Still verify the shape and the link.",
        "Scores between 0.50 and 0.80 — uncertain. Inspect carefully. Some "
        "are real contrails, others are false positives.",
        "Scores around 0.50 — very suspicious. Often sun artifacts or thin "
        "clouds. Delete if no real contrail is visible.",
    ])

    add_tip(doc,
        "Use the score to prioritise your attention. Focus extra scrutiny "
        "on low-score detections — they are the most likely to be wrong."
    )

    # ── 3.5 Out-of-scope contrails ──
    doc.add_heading(
        "3.5  Out-of-scope contrails (no Prompt = do not label)", level=2
    )
    doc.add_paragraph(
        "You will sometimes see contrails in the sky that do not correspond "
        "to any Prompt. These are contrails produced by flights that passed "
        "outside the camera's field of view, or by flights for which we do "
        "not have trajectory data. Because we cannot identify which flight "
        "created them, no Prompt exists for these contrails."
    )

    add_important(doc,
        "If a contrail has no matching Prompt, it is out of scope. Do NOT "
        "label it. Do NOT draw a new Contrail polygon for it. Only label "
        "contrails that can be linked to a Prompt — that is the entire "
        "purpose of this campaign."
    )

    doc.add_paragraph(
        "How to recognise an out-of-scope contrail:"
    )
    add_bullets(doc, [
        "It is visible in the image but does not align with any orange "
        "Prompt polygon.",
        "It may have been present since the start of the sequence (an old "
        "contrail already in the sky before the 2-hour window began).",
        "It may be produced by a flight that crossed the sky outside the "
        "camera's field of view — only the drifted contrail enters the "
        "image, without any corresponding Prompt.",
    ])

    doc.add_paragraph(
        "These contrails are real, but because we cannot associate them "
        "with a specific flight, they are scientifically unusable for this "
        "campaign. Ignore them completely."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. HOW THE MODEL WORKS
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("4  How the model works", level=1)

    doc.add_paragraph(
        "Understanding how the model works helps you spot its mistakes. "
        "The process has two steps:"
    )

    doc.add_heading("Step 1: Prompts are generated from flight data", level=2)
    doc.add_paragraph(
        "For each flight that passed over the camera, the system projects "
        "the aircraft's last 5 minutes of trajectory onto the camera image. "
        "Wind data is used to shift the projection (advection), so the Prompt "
        "shows where the contrail would have drifted to — not just where the "
        "aircraft flew. This creates an orange Prompt polygon. Different flights "
        "produce different Prompts, each appearing on the frames when that "
        "flight was overhead."
    )

    doc.add_heading("Step 2: The model detects and tracks contrails", level=2)
    doc.add_paragraph(
        "The SAM2 model looks at each Prompt region and decides whether a "
        "contrail is present. If it finds one, it draws a red multi-polygon "
        "around it and assigns a confidence score. Crucially, the model also "
        "tracks each contrail across consecutive frames — so a single "
        "Contrail instance spans many frames of the video, following the "
        "contrail as it evolves."
    )

    doc.add_paragraph("")
    add_placeholder(doc, "Frame with several Prompts (orange) and Contrails (red)")
    add_caption(doc,
        "Figure 2 — Example frame showing several Prompts (orange polygons) "
        "and predicted Contrails (red polygons)."
    )

    doc.add_heading("How contrails evolve over time", level=2)
    doc.add_paragraph(
        "A contrail starts as a thin white line behind the aircraft. Over "
        "minutes to hours it spreads wider, may develop irregular edges, and "
        "eventually merges with natural cirrus clouds. The model tracks each "
        "contrail across many frames of the video. The figure below shows how "
        "a sequence evolves over 30 minutes:"
    )

    doc.add_paragraph("")
    add_placeholder(doc,
        "Sequence evolution: 4 panels showing frames 1, 5, 10, 15 "
        "— Prompts appear/disappear, Contrails emerge, spread, and dissipate"
    )
    add_caption(doc,
        "Figure 3 — Sequence evolution during 15 frames (30 minutes). Notice "
        "how Prompts appear and disappear as different flights pass overhead, "
        "and how Contrails emerge, spread, change shape, and eventually "
        "dissipate."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. USING TEMPORAL NAVIGATION
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("5  Using temporal navigation", level=1)

    doc.add_paragraph(
        "Because you are reviewing video sequences, navigating forward and "
        "backward through the frames is by far your most effective tool. "
        "A single frame can be ambiguous — is that a contrail or a wisp of "
        "cirrus? Is that polygon on a real feature or on sun glare? Almost "
        "every ambiguity is resolved by watching how the feature behaves "
        "over time."
    )

    doc.add_heading("What temporal navigation reveals", level=2)

    p = doc.add_paragraph()
    run = p.add_run("1. Is it a real contrail?  ")
    run.bold = True
    p.add_run(
        "Real contrails appear when a flight passes overhead, persist for "
        "many frames (minutes to hours), and drift gradually with the wind. "
        "Sun artifacts, by contrast, appear and vanish abruptly as the sun "
        "angle changes, or stay fixed relative to the sun position rather "
        "than drifting with the wind. By scrubbing a few frames forward or "
        "backward, you can immediately tell the difference."
    )

    p = doc.add_paragraph()
    run = p.add_run("2. Does the Contrail belong to this Prompt?  ")
    run.bold = True
    p.add_run(
        "If a Contrail is correctly linked to a Prompt, you should see it "
        "emerge around the same time and location as the Prompt appears, "
        "and the two should stay aligned in direction across frames. "
        "If the Contrail was already present before the Prompt appeared, "
        "or if it moves in a different direction, the link is probably wrong."
    )

    p = doc.add_paragraph()
    run = p.add_run("3. Is the polygon shape correct?  ")
    run.bold = True
    p.add_run(
        "Watching the contrail evolve frame by frame makes it easy to judge "
        "whether the polygon follows the real boundary. A polygon that looks "
        "reasonable in one frame but jumps erratically to a different region "
        "in the next is a sign that the model made an error on some frames."
    )

    p = doc.add_paragraph()
    run = p.add_run("4. When should labelling stop?  ")
    run.bold = True
    p.add_run(
        "Stepping forward through frames lets you see exactly when a contrail "
        "becomes too diffuse to label. Once you can no longer distinguish it "
        "from the background sky or from natural clouds, that is where the "
        "Contrail annotation should end."
    )

    add_important(doc,
        "Never judge a prediction from a single frame alone. Always scrub "
        "through at least a few neighbouring frames before deciding to "
        "accept, correct, or delete a Contrail polygon."
    )

    add_tip(doc,
        "In Encord, use the left/right arrow keys to step through frames "
        "one at a time. This is the fastest way to review a Contrail's "
        "temporal behaviour."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. YOUR TASKS AS A LABELLER
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("6  Your tasks as a labeller (Prompt by Prompt)", level=1)

    doc.add_paragraph(
        "The review should be done Prompt by Prompt. For each 2-hour video "
        "sequence, go through the list of Prompts one at a time. For each "
        "Prompt, navigate through the frames where it appears and perform "
        "the checks described below. This ensures no flight is overlooked "
        "and every prediction is verified in context."
    )

    add_important(doc,
        "Work Prompt by Prompt, not Contrail by Contrail. Start from the "
        "Prompt (the flight), then check what the model predicted for that "
        "flight. Use temporal navigation to follow the Prompt and its "
        "associated Contrail across frames."
    )

    # ── 6.1 ──
    doc.add_heading(
        "6.1  For each Prompt: check if a real contrail is visible", level=2
    )

    doc.add_paragraph(
        "Select a Prompt (orange polygon) and navigate through the frames "
        "where it appears. Ask yourself: is there a visible contrail in the "
        "image that is aligned with this Prompt?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Case A — A Contrail prediction exists and is correct:  ")
    run.bold = True
    p.add_run(
        "The model has drawn a red polygon that matches the visible contrail, "
        "and it is linked to this Prompt. Step through a few frames to verify "
        "that the polygon tracks the contrail consistently. If the shape is "
        "reasonable across frames, move on."
    )

    p = doc.add_paragraph()
    run = p.add_run("Case B — A Contrail prediction exists but is wrong:  ")
    run.bold = True
    p.add_run(
        "The red polygon may be on a sun artifact, a cloud, or simply in the "
        "wrong place. Use temporal navigation to confirm: does the feature "
        "persist and drift like a real contrail, or does it appear/vanish "
        "abruptly? If it is a false positive, delete the red Contrail polygon "
        "(see Section 6.4). If it belongs to a different flight, re-link it "
        "to the correct Prompt."
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Case C — No Contrail prediction, but a contrail is visible:  "
    )
    run.bold = True
    p.add_run(
        "The model missed the contrail. Draw a new Contrail polygon (red) "
        "around it, set its class to \"Contrail\", and link it to this Prompt "
        "via the flight #relation attribute. Navigate through frames to draw "
        "the polygon on each frame where the contrail is visible."
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Case D — No Contrail prediction and no contrail visible:  "
    )
    run.bold = True
    p.add_run(
        "This is normal — not every flight produces a visible contrail. "
        "Do nothing and move to the next Prompt."
    )

    add_tip(doc,
        "You may notice contrails in the sky that do not correspond to any "
        "Prompt. These are out-of-scope contrails from flights outside the "
        "camera's range — ignore them completely (see Section 3.5)."
    )

    doc.add_paragraph("")
    add_placeholder(doc,
        "Two panels: (left) a young contrail overlapping its Prompt; "
        "(right) the same contrail later, now extended beyond the Prompt "
        "but still aligned"
    )
    add_caption(doc,
        "Figure 4 — A Contrail at an early stage overlapping the Prompt "
        "(left), and the same Contrail after some time — no longer "
        "overlapping but still aligned with the Prompt direction (right)."
    )

    doc.add_paragraph("")
    add_placeholder(doc, "A Prompt with no associated Contrail (Case D)")
    add_caption(doc,
        "Figure 5 — A Prompt without a Contrail. No visible contrail is "
        "present — this flight did not produce one. This is normal; do not "
        "create a Contrail annotation."
    )

    # ── 6.2 ──
    doc.add_heading("6.2  Verify the Contrail-to-Prompt link", level=2)

    doc.add_paragraph(
        "When a Contrail prediction exists for a Prompt, verify that the "
        "model linked them correctly:"
    )

    p = doc.add_paragraph()
    run = p.add_run("How to check:  ")
    run.bold = True
    p.add_run(
        "Click on the red Contrail polygon. In the attributes panel, check "
        "the flight #relation field — it should point to the correct Prompt. "
        "Then visually verify across several frames: is the Contrail aligned "
        "with that Prompt (same direction)? Did it appear around the same "
        "time as the Prompt?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Signs of a wrong link:  ")
    run.bold = True
    p.add_run(
        "The Contrail runs in a completely different direction than the "
        "Prompt. The Contrail was already present before the Prompt appeared. "
        "The Contrail is far away with no spatial relationship to the "
        "Prompt's trajectory."
    )

    p = doc.add_paragraph()
    run = p.add_run("How to fix:  ")
    run.bold = True
    p.add_run(
        "Change the flight #relation to point to the correct Prompt. "
        "Look at nearby Prompts to find the right match — the correct "
        "Prompt will be the one whose direction matches the Contrail."
    )

    # ── 6.3 ──
    doc.add_heading("6.3  Fix polygon shapes", level=2)

    doc.add_paragraph(
        "Sometimes the model detects a contrail correctly, but the polygon "
        "shape is not accurate — it may be too large, too small, or poorly "
        "aligned with the contrail boundary."
    )

    p = doc.add_paragraph()
    run = p.add_run("How to fix:  ")
    run.bold = True
    p.add_run(
        "Click on the polygon and drag its vertices to better follow the "
        "contrail boundary. Step through neighbouring frames to check that "
        "the shape is consistent. Make sure the polygon covers the full "
        "visible extent of the contrail on each frame."
    )

    add_bullets(doc, [
        "If the contrail is fragmented or partially hidden by clouds, it is "
        "OK for the polygon to cover only the visible parts.",
        "Do NOT label features that have spread so much they look like "
        "natural clouds. Only label contrails that are still clearly "
        "identifiable — relatively linear and distinct from the background.",
    ])

    # ── 6.4 ──
    doc.add_heading("6.4  Delete false positives (sun artifacts)", level=2)

    doc.add_paragraph(
        "The model's most common mistake is confusing sun glare with "
        "contrails. This happens in morning and evening frames when the "
        "sun is low on the horizon."
    )

    p = doc.add_paragraph()
    run = p.add_run("How to recognise a sun artifact:  ")
    run.bold = True

    add_bullets(doc, [
        "It appears near a bright, glowing area of the image (sun glare).",
        "It has a low confidence score (typically 0.50 to 0.55).",
        "It does not look like a thin line — often a blob or irregular patch.",
        "It does not drift with the wind like a real contrail. Instead, it "
        "stays fixed relative to the sun's position in the image.",
        "It appears and disappears abruptly as the sun angle changes "
        "between frames.",
    ])

    add_important(doc,
        "Temporal navigation is the clearest way to identify sun artifacts. "
        "Step forward and backward a few frames: if the feature does not "
        "persist and drift like a real contrail, it is a sun artifact."
    )

    p = doc.add_paragraph()
    run = p.add_run("What to do:  ")
    run.bold = True
    p.add_run(
        "DELETE the red Contrail polygon. Do NOT delete the orange Prompt "
        "— the Prompt is just the flight trajectory and must always remain."
    )

    doc.add_paragraph("")
    add_placeholder(doc, "Sun glare false positive example")
    add_caption(doc,
        "Figure 6 — Example of a sun artifact falsely detected as a Contrail. "
        "Notice the bright glare and the low score. This should be deleted."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. COMMON MISTAKES TO AVOID
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("7  Common mistakes to avoid", level=1)

    doc.add_heading("7.1  Judging from a single frame", level=2)
    doc.add_paragraph(
        "The most common labeller mistake is making decisions based on one "
        "frame. Always navigate through several frames before accepting, "
        "correcting, or deleting a prediction. The temporal context makes "
        "almost every ambiguous case clear."
    )

    doc.add_heading("7.2  Confusing cirrus clouds with contrails", level=2)
    doc.add_paragraph(
        "Cirrus clouds are thin, wispy natural clouds at high altitude. "
        "They can look similar to old contrails. Here is how to tell them "
        "apart:"
    )
    add_styled_table(doc,
        headers=["Feature", "Contrail", "Cirrus cloud"],
        rows=[
            ["Shape",
             "Linear or slightly curved; can be elongated or rounded/irregular",
             "Wispy, curved, very irregular"],
            ["Alignment",
             "Aligned with a Prompt (same direction)",
             "Not aligned with any Prompt"],
            ["Temporal behaviour",
             "Appears when a flight passes; drifts with wind; gradually spreads",
             "Develops slowly; part of a broader cloud layer"],
            ["Context",
             "Often part of a system of parallel lines",
             "Part of a diffuse cloud field"],
        ],
        col_widths=[2.5, 5.5, 5.5],
    )

    doc.add_paragraph("")
    add_tip(doc,
        "If a feature does not align with any Prompt, it is most likely a "
        "natural cloud. Only label it as a Contrail if it clearly follows "
        "a Prompt's direction and appeared when that flight was overhead."
    )

    doc.add_heading(
        "7.3  Expecting Prompts and Contrails to overlap perfectly", level=2
    )
    doc.add_paragraph(
        "Prompts only cover the last 5 minutes of the flight trajectory. "
        "A contrail that has survived longer than 5 minutes will extend "
        "beyond its Prompt. This is normal. The key check is alignment "
        "(same direction), not pixel-level overlap. Use temporal navigation "
        "to watch the Prompt appear and the Contrail extend beyond it "
        "over time."
    )

    doc.add_heading("7.4  Labelling contrails that are too old", level=2)
    doc.add_paragraph(
        "Contrails spread over time and eventually merge with natural "
        "clouds. Only label contrails that are still clearly identifiable "
        "— relatively linear and distinct from the background. Step forward "
        "through frames: once you can no longer distinguish the contrail "
        "from the surrounding sky, that is where the annotation should stop."
    )

    doc.add_heading("7.5  Forgetting to set the flight #relation", level=2)
    doc.add_paragraph(
        "If you draw a new Contrail polygon, you MUST set its flight "
        "#relation attribute to link it to the correct Prompt. A Contrail "
        "without a link is unusable for our analysis."
    )

    doc.add_heading("7.6  Labelling contrails without a Prompt", level=2)
    doc.add_paragraph(
        "You may see contrails in the sky that have no corresponding Prompt. "
        "These were produced by flights outside the camera's field of view "
        "or outside the flight data coverage. Because they cannot be linked "
        "to a flight, they are out of scope and must never be labelled. "
        "Only label contrails that can be associated with a Prompt."
    )

    doc.add_heading("7.7  Deleting Prompts", level=2)
    doc.add_paragraph(
        "Never delete a Prompt (orange polygon). Prompts represent flight "
        "trajectories and are part of the input data. Even if no contrail "
        "is visible for a given Prompt, the Prompt must stay."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. STEP-BY-STEP ENCORD WORKFLOW
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("8  Step-by-step Encord workflow", level=1)

    doc.add_paragraph(
        "Follow this procedure for every video sequence assigned to you:"
    )

    steps = [
        ("Open the video sequence",
         "In the Encord project, click on a sequence to open it in the "
         "editor. You will see the first frame with pre-annotated Prompt "
         "(orange) and Contrail (red) polygons."),

        ("Get an overview of the sequence",
         "Before reviewing individual Prompts, scrub through the full "
         "sequence using the timeline bar or arrow keys. This gives you a "
         "sense of how many contrails are present, where the sun glare is, "
         "and which parts of the video are busiest."),

        ("Identify all Prompts",
         "In the objects panel, note the list of Prompts. Each Prompt "
         "represents one flight. You will review them one by one."),

        ("Select a Prompt and navigate its frames",
         "Click on an orange Prompt polygon to highlight it. Note its "
         "flight_id (shown in the attributes panel). Use the arrow keys "
         "to step through the frames where this Prompt appears. Watch how "
         "it enters and exits the frame."),

        ("Check for a linked Contrail",
         "Look for a red Contrail polygon that is aligned with this Prompt. "
         "If one exists, click on it and verify that its flight #relation "
         "points to this Prompt. Check its score attribute. Then step "
         "through frames to verify: does the red polygon consistently "
         "track a visible contrail?"),

        ("Assess the prediction using temporal context",
         "Is the red polygon on a real contrail? Does it persist and drift "
         "with the wind across frames, or does it appear/vanish abruptly? "
         "Is the shape reasonable across frames? "
         "If correct: move on. "
         "If it is a false positive (sun artifact, cloud): delete the red "
         "polygon. "
         "If the shape is wrong on some frames: adjust the vertices."),

        ("Check for missed contrails",
         "If no red polygon exists but you can see a visible contrail "
         "aligned with this Prompt across several frames, draw a new "
         "Contrail polygon. Set its class to \"Contrail\" and link it to "
         "this Prompt via flight #relation."),

        ("Move to the next Prompt",
         "Repeat steps 4–7 for every Prompt in the sequence."),

        ("Final scan",
         "Once all Prompts are reviewed, do a quick forward scan of the "
         "entire sequence. This catches anything you might have missed, "
         "such as a contrail that is visible but has no Prompt (rare but "
         "possible if the flight data was incomplete)."),

        ("Submit",
         "When you have reviewed all Prompts and completed the final scan, "
         "click Submit."),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i}: {title}")
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(desc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. QUICK REFERENCE CHECKLIST
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("9  Quick reference checklist", level=1)

    doc.add_paragraph(
        "Use this checklist for every video sequence you review:"
    )

    checklist = [
        "Reviewed every Prompt (orange) — checked whether a contrail "
        "is visible for that flight.",
        "Used temporal navigation (arrow keys) to verify predictions "
        "across multiple frames, not just one.",
        "Every Contrail (red) is linked to the correct Prompt "
        "(flight #relation is set).",
        "No Contrail polygons remain on sun artifacts or non-contrail "
        "features.",
        "Polygon shapes follow the contrail boundary reasonably well "
        "across frames.",
        "Missed contrails have been drawn and linked to the right Prompt.",
        "No Prompts have been deleted or modified.",
        "No out-of-scope contrails labelled (contrails without a matching "
        "Prompt are ignored).",
        "Contrails that have spread into diffuse clouds are not labelled.",
        "Final full-sequence scan completed.",
    ]

    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")

    add_styled_table(doc,
        headers=["Colour", "Object", "Action"],
        rows=[
            ["Orange", "Prompt",
             "Never delete. Never edit. This is input data."],
            ["Red", "Contrail",
             "Verify, correct shape, fix link, or delete if false positive."],
        ],
        col_widths=[2, 3, 10],
    )

    doc.add_paragraph("")
    add_tip(doc,
        "When in doubt, use the video. Step forward and backward through "
        "frames. Real contrails persist and drift; artifacts appear and "
        "vanish. The temporal information is your most reliable signal."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.save(str(OUT_PATH))
    print(f"Document saved to {OUT_PATH}")


if __name__ == "__main__":
    build_document()
